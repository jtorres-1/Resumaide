import os
from flask import Flask, render_template, request, send_file, session, redirect, url_for
import openai
from PyPDF2 import PdfReader
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_limiter.errors import RateLimitExceeded

# Load environment variables
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "supersecretkey")  # needed for session

# --- Add API rate limiting ---
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=[]
)

@app.route("/", methods=["GET", "POST"])
@limiter.limit("3 per day", methods=["POST"])
def index():
    if request.method == "POST":
        resume_text = ""

        # Handle PDF upload
        if "resume_file" in request.files and request.files["resume_file"].filename != "":
            pdf_file = request.files["resume_file"]
            try:
                reader = PdfReader(pdf_file)
                resume_text = " ".join([page.extract_text() or "" for page in reader.pages])
            except Exception as e:
                return render_template("index.html", error=f"Failed to read PDF: {e}")

        # Handle pasted text
        if not resume_text and request.form.get("resume_text"):
            resume_text = request.form["resume_text"]

        if not resume_text.strip():
            return render_template("index.html", error="No resume provided.")

        # Send to OpenAI
        prompt = f"""
        Rewrite this resume to be clear, keyword-rich, and formatted to pass ATS (Applicant Tracking Systems).
        Only return the improved resume, nothing else. Make it the best resume ever written for this user.
        Make them seem like a professional.

        Resume:
        {resume_text}
        """

        try:
            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You optimize resumes for ATS parsing."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=1200,
            )
            optimized_resume = response.choices[0].message.content.strip()

            # Store in session (per-user, isolated)
            session["optimized_resume"] = optimized_resume

            return render_template("result.html", optimized_resume=optimized_resume)

        except Exception as e:
            return render_template("index.html", error=f"Error: {e}")

    return render_template("index.html")


@app.route("/download/txt", methods=["POST"])
def download_txt():
    optimized_resume = session.get("optimized_resume", "")
    if not optimized_resume:
        return redirect(url_for("index"))

    buffer = BytesIO()
    buffer.write(optimized_resume.encode("utf-8"))
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="optimized_resume.txt", mimetype="text/plain")


@app.route("/download/docx", methods=["POST"])
def download_docx():
    optimized_resume = session.get("optimized_resume", "")
    if not optimized_resume:
        return redirect(url_for("index"))

    doc = Document()
    for line in optimized_resume.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")  # preserve spacing
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="optimized_resume.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.errorhandler(RateLimitExceeded)
def ratelimit_handler(e):
    return render_template("429.html"), 429


if __name__ == "__main__":
    app.run(debug=True)
